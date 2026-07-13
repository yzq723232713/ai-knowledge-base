"""文档加载器：PDF / DOCX / TXT / CSV → Document"""
from pathlib import Path
import csv
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from src.models import Document

import logging
logger = logging.getLogger(__name__)

class DocumentLoader:
    """根据扩展名自动路由到对应的加载器"""

    SUPPORTED = {".pdf", ".docx", ".txt", ".csv"}

    @staticmethod
    def load(file_path: str) -> Document:
        """
        加载文档，损坏的文件不崩溃，log 警告并返回空 Document。
        调用方检查 Document.content 是否为空即可判断加载是否成功。
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix not in DocumentLoader.SUPPORTED:
            logger.warning(f"跳过不支持格式: {file_path}")
            return Document(content="", source=path.name, file_type="unknown")

        try:
            if suffix == ".pdf":
                return DocumentLoader._load_pdf(path)
            elif suffix == ".docx":
                return DocumentLoader._load_docx(path)
            elif suffix == ".txt":
                return DocumentLoader._load_txt(path)
            else:
                return DocumentLoader._load_csv(path)
        except Exception as e:
            logger.warning(f"加载失败 {path.name}: {e}")
            return Document(content="", source=path.name, file_type=suffix.replace(".", ""))

    @staticmethod
    def _load_pdf(path: Path) -> Document:
        doc = fitz.open(path)
        full_text = []
        for page in doc:
            full_text.append(page.get_text())
        page_count = doc.page_count
        doc.close()
        return Document(
            content="\n".join(full_text).strip(),
            source=path.name,
            file_type="pdf",
            page_count=page_count,
        )

    @staticmethod
    def _load_docx(path: Path) -> Document:
        doc = DocxDocument(path)
        content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return Document(content=content, source=path.name, file_type="docx")

    @staticmethod
    def _load_txt(path: Path) -> Document:
        content = path.read_text(encoding="utf-8").strip()
        return Document(content=content, source=path.name, file_type="txt")

    @staticmethod
    def _load_csv(path: Path) -> Document:
        with open(path, encoding="utf-8-sig") as f:
            rows = list(csv.DictReader(f))
        if not rows:
            return Document(content="", source=path.name, file_type="csv")
        headers = list(rows[0].keys())
        lines = [" | ".join(headers)]
        lines.append(" | ".join(["---"] * len(headers)))
        for row in rows:
            lines.append(" | ".join(str(row[h]) for h in headers))
        return Document(content="\n".join(lines), source=path.name, file_type="csv")